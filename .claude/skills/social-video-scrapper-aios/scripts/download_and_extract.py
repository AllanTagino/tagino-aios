"""
Social Video Scrapper helper script.
Downloads videos via yt-dlp, extracts transcripts (subtitles first, Whisper fallback),
cleans them, and outputs .txt (agents) and .docx / .pdf (humans).
"""

import argparse
import re
import subprocess
import sys
from pathlib import Path


# ---------- subprocess helpers ----------

def run(cmd: list[str], label: str, allow_fail: bool = False) -> tuple[bool, str]:
    print(f"[social-video-scrapper] {label}...", flush=True)
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0 and not allow_fail:
        print(f"[ERROR] {label} failed:\n{result.stderr}", file=sys.stderr)
        sys.exit(result.returncode)
    return result.returncode == 0, result.stdout


# ---------- text cleanup ----------

TEXT_FIXES = {
    "clawed": "Claude",
    "claw.ai": "claude.ai",
    "Claw Code": "Claude Code",
    "muchneeded": "much-needed",
    "MGB": "MB",
    "NodeJS": "Node.js",
    "Google doc": "Google Doc",
    "MD extension": ".md extension",
    "MD file": ".md file",
    ".MD": ".md",
    "open- source": "open-source",
    "well- formatted": "well-formatted",
}


def clean_text(text: str) -> str:
    for bad, good in TEXT_FIXES.items():
        text = text.replace(bad, good)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"([.,!?])([A-Za-z])", r"\1 \2", text)
    return text


def clean_vtt(vtt_path: Path) -> str:
    content = vtt_path.read_text(encoding="utf-8")
    lines = content.split("\n")
    clean: list[str] = []
    seen: set[str] = set()
    for line in lines:
        line = line.strip()
        if (
            not line
            or line.startswith("WEBVTT")
            or line.startswith("Kind:")
            or line.startswith("Language:")
            or "-->" in line
            or line.isdigit()
        ):
            continue
        line = re.sub(r"<[^>]+>", "", line).strip()
        if line and line not in seen:
            seen.add(line)
            clean.append(line)
    return clean_text(" ".join(clean))


def split_paragraphs(text: str) -> list[str]:
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
    paras, buf = [], []
    for i, s in enumerate(sentences):
        buf.append(s)
        next_s = sentences[i + 1] if i + 1 < len(sentences) else ""
        if len(buf) >= 4 and (
            len(buf) >= 5 or re.match(r"^(Now|So|Next|Then|Once|Finally|But|However)\b", next_s)
        ):
            paras.append(" ".join(buf))
            buf = []
    if buf:
        paras.append(" ".join(buf))
    return paras


# ---------- transcription ----------

def extract_subtitles(url: str, output_stem: Path, lang: str) -> str | None:
    vtt_path = output_stem.with_suffix(f".{lang}.vtt")
    ok, _ = run(
        ["yt-dlp", "--write-auto-sub", "--sub-lang", lang, "--skip-download",
         "-o", str(output_stem), url],
        f"Extracting {lang} subtitles",
        allow_fail=True,
    )
    if ok and vtt_path.exists():
        print(f"[social-video-scrapper] Subtitles saved to: {vtt_path}")
        return clean_vtt(vtt_path)
    return None


def extract_whisper(video_path: Path) -> str:
    try:
        import whisper  # type: ignore
    except ImportError:
        print("[social-video-scrapper] Whisper not installed. Run: pip install openai-whisper", file=sys.stderr)
        sys.exit(1)
    print("[social-video-scrapper] Running Whisper transcription...", flush=True)
    model = whisper.load_model("base")
    return clean_text(model.transcribe(str(video_path))["text"])


# ---------- docx / pdf output ----------

def write_docx(docx_path: Path, title: str, transcript: str, summary_bullets: list[str] | None = None) -> None:
    from docx import Document  # type: ignore
    from docx.shared import Pt

    doc = Document()

    t = doc.add_paragraph()
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"

    if summary_bullets:
        h = doc.add_paragraph()
        hr = h.add_run("Summary")
        hr.bold = True
        hr.font.size = Pt(16)
        hr.font.name = "Calibri"
        for b in summary_bullets:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(b)
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    h2 = doc.add_paragraph()
    h2r = h2.add_run("Full Transcript")
    h2r.bold = True
    h2r.font.size = Pt(16)
    h2r.font.name = "Calibri"

    for para in split_paragraphs(transcript):
        p = doc.add_paragraph()
        run = p.add_run(para)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.3

    doc.save(docx_path)
    print(f"[social-video-scrapper] DOCX saved: {docx_path}")


SOFFICE_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/usr/bin/soffice",
    "/usr/local/bin/soffice",
]


def write_pdf(docx_path: Path, pdf_path: Path) -> None:
    print(f"[social-video-scrapper] Converting to PDF...", flush=True)

    # Try LibreOffice first (always available, no COM dependency)
    soffice = next((p for p in SOFFICE_PATHS if Path(p).exists()), None)
    if soffice:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            capture_output=True, text=True,
        )
        if result.returncode == 0:
            # LibreOffice names the file <stem>.pdf in outdir
            lo_out = docx_path.with_suffix(".pdf")
            if lo_out != pdf_path and lo_out.exists():
                lo_out.rename(pdf_path)
            print(f"[social-video-scrapper] PDF saved: {pdf_path}")
            return
        print(f"[social-video-scrapper] LibreOffice failed: {result.stderr.strip()}", file=sys.stderr)

    # Fallback: docx2pdf (requires MS Word COM on Windows)
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(docx_path), str(pdf_path))
        print(f"[social-video-scrapper] PDF saved: {pdf_path}")
    except ImportError:
        print("[social-video-scrapper] PDF skipped — install LibreOffice or run: pip install docx2pdf", file=sys.stderr)
    except Exception as e:
        print(
            f"[social-video-scrapper] PDF conversion failed ({e.__class__.__name__}). "
            f"DOCX is available at: {docx_path}",
            file=sys.stderr,
        )


# ---------- main ----------

def main():
    parser = argparse.ArgumentParser(description="Download and optionally transcribe/analyze a social video.")
    parser.add_argument("--url", required=True, help="YouTube or Instagram URL")
    parser.add_argument("--output", default=str(Path.home() / "Downloads" / "video.mp4"),
                        help="Output video file path (default: ~/Downloads/video.mp4)")
    parser.add_argument("--transcribe", action="store_true", help="Extract transcript")
    parser.add_argument("--whisper", action="store_true", help="Force Whisper transcription")
    parser.add_argument("--lang", default="en", help="Subtitle language code (default: en)")
    parser.add_argument("--docx", action="store_true", help="Also write a polished .docx")
    parser.add_argument("--pdf", action="store_true", help="Also convert .docx to .pdf (implies --docx)")
    parser.add_argument("--title", default="Video Summary", help="Title for the DOCX/PDF")
    args = parser.parse_args()

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    run(["yt-dlp", "-o", str(output_path), args.url], f"Downloading {args.url}")
    print(f"[social-video-scrapper] Downloaded to: {output_path}")

    need_transcript = args.transcribe or args.docx or args.pdf
    if not need_transcript:
        return

    if args.whisper:
        transcript = extract_whisper(output_path)
    else:
        transcript = extract_subtitles(args.url, output_path.with_suffix(""), args.lang)
        if transcript is None:
            print("[social-video-scrapper] No subtitles found — falling back to Whisper.")
            transcript = extract_whisper(output_path)

    # Always save the clean .txt for agents
    txt_path = output_path.with_suffix(".txt")
    txt_path.write_text(transcript, encoding="utf-8")
    print(f"[social-video-scrapper] Transcript saved: {txt_path}")

    if args.docx or args.pdf:
        docx_path = output_path.with_suffix(".docx")
        write_docx(docx_path, args.title, transcript)
        if args.pdf:
            write_pdf(docx_path, output_path.with_suffix(".pdf"))


if __name__ == "__main__":
    main()
